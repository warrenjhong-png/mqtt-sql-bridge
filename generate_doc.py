from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── 頁面設定 ──────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.27)
section.page_height = Inches(11.69)
section.left_margin = section.right_margin = Cm(2.5)
section.top_margin  = section.bottom_margin = Cm(2.5)

# ── 全域字型 ──────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = '微軟正黑體'
style.font.size = Pt(11)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def h3(text):
    return doc.add_heading(text, level=3)

def body(text):
    return doc.add_paragraph(text)

def bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x80)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    p._p.get_or_add_pPr().append(shd)
    return p

def table2(rows, headers=None):
    cols = 2
    t = doc.add_table(rows=len(rows) + (1 if headers else 0), cols=cols)
    t.style = 'Table Grid'
    offset = 0
    if headers:
        for i, h in enumerate(headers):
            cell = t.rows[0].cells[i]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
        offset = 1
    for r, (a, b) in enumerate(rows):
        t.rows[r + offset].cells[0].text = str(a)
        t.rows[r + offset].cells[1].text = str(b)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════
# 封面
# ══════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('MQTT → SQL Server Bridge')
run.font.size = Pt(26)
run.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run('系統交接文件').font.size = Pt(16)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('版本：v1.0　　日期：2026-06-07').font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 1. 系統概述
# ══════════════════════════════════════════════════════════
h1('1. 系統概述')
body('本系統為一個 MQTT 到 SQL Server 的資料橋接服務，'
     '負責即時接收工廠設備（空壓機）透過 MQTT 協定發布的感測器資料，'
     '並寫入 Microsoft SQL Server 資料庫，同時在本地備份原始 JSON 資料。')

doc.add_paragraph()
h2('1.1 系統架構圖')
body('  MQTT Broker（109.123.238.225:1883）')
body('       │  訂閱 topic（QoS 1）')
body('       ▼')
body('  MQTTReceiver')
body('       │  解析 JSON payload')
body('       ├──→ RawJsonWriter（本地備份 .jsonl）')
body('       └──→ DataAggregator（pass-through 或降頻平均）')
body('                  │')
body('                  ▼')
body('           Queue（in-memory buffer）')
body('                  │')
body('                  ▼')
body('           DBWriter → SQL Server')

doc.add_paragraph()
h2('1.2 部署環境')
table2([
    ('作業系統', 'Windows 11 + WSL2 Ubuntu 22.04'),
    ('Docker', 'Docker Engine（WSL2 內）'),
    ('容器名稱', 'mqtt_sql_bridge'),
    ('專案路徑（Windows）', r'C:\Users\autolab\Documents\code\mqtt-sql-bridge'),
    ('專案路徑（WSL2）', '~/mqtt-sql-bridge'),
], headers=['項目', '說明'])

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 2. 系統邏輯與機制
# ══════════════════════════════════════════════════════════
h1('2. 系統邏輯與機制')

h2('2.1 MQTT 接收（MQTTReceiver）')
bullet('使用 paho-mqtt 連線至 broker，斷線後自動重連')
bullet('依 config.yaml 訂閱多個 topic，每個 topic 對應一個 DB table')
bullet('收到訊息後：')
bullet('先備份原始 JSON 至本地 raw_json/ 目錄', level=1)
bullet('再將 payload 傳入 DataAggregator', level=1)
bullet('空 payload 會被跳過並記錄警告')

doc.add_paragraph()
h2('2.2 資料緩衝（Queue）')
bullet('MQTTReceiver 與 DBWriter 之間有一個 in-memory queue 做緩衝')
bullet('預設上限 10,000 筆（可在 config.yaml 調整 queue_maxsize）')
bullet('container 重啟後 queue 清空，不會補寫舊資料')

doc.add_paragraph()
h2('2.3 資料庫寫入（DBWriter）')
bullet('從 queue 取出資料，依 field_mapping.yaml 組合 INSERT SQL')
bullet('INSERT 欄位：TIMETAG（payload 的 Timestamp 欄位）、TIME01（GETDATE()）、FIELD_1～FIELD_17')
bullet('若 payload 沒有 Timestamp，TIMETAG 自動 fallback 為 GETDATE()')
bullet('Timestamp 字串（ISO 8601）會在 Python 端解析後再傳入，避免 SQL Server 轉換錯誤')
bullet('DB 斷線後自動重連（reconnect_delay 秒後重試）')
bullet('每 5 分鐘發送 SELECT 1 心跳，防止防火牆切斷閒置連線')
bullet('INSERT 失敗時記錄 ERROR log（資料已備份於 raw_json）')

doc.add_paragraph()
h2('2.4 降頻模式（DataAggregator）')
bullet('config.yaml 的 sampling_interval_seconds = 0：pass-through，每筆立即寫入')
bullet('設定為 N（秒）：每 N 秒將收到的所有資料平均為一筆後寫入')

doc.add_paragraph()
h2('2.5 本地備份（RawJsonWriter）')
bullet('每筆收到的 MQTT 訊息都會以 JSONL 格式備份')
bullet('按天分檔：raw_json/YYYY-MM-DD.jsonl')
bullet('無論 DB 寫入成功或失敗，備份都會執行')
bullet('備份檔不會自動清除，需手動管理磁碟空間')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 3. 檔案結構
# ══════════════════════════════════════════════════════════
h1('3. 檔案結構')
code_block(
    'mqtt-sql-bridge/\n'
    '├── main.py              # 程式進入點，啟動所有元件\n'
    '├── mqtt_receiver.py     # MQTT 連線與訊息接收\n'
    '├── db_writer.py         # DB 連線與資料寫入\n'
    '├── aggregator.py        # 降頻平均邏輯\n'
    '├── raw_json_writer.py   # 本地 JSON 備份\n'
    '├── config_loader.py     # 讀取 config.yaml\n'
    '├── logger.py            # 日誌設定\n'
    '├── simulator.py         # 測試用 MQTT 資料發布工具\n'
    '├── config.yaml          # 主要設定檔（已加入 .gitignore）\n'
    '├── field_mapping.yaml   # MQTT payload 欄位對應 DB 欄位\n'
    '├── requirements.txt     # Python 套件清單\n'
    '├── Dockerfile           # Docker image 建置設定\n'
    '├── docker-compose.yml   # Container 啟動設定\n'
    '├── logs/                # 日誌目錄（已加入 .gitignore）\n'
    '└── raw_json/            # 備份目錄（已加入 .gitignore）'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 4. 設定檔說明
# ══════════════════════════════════════════════════════════
h1('4. 設定檔說明')

h2('4.1 config.yaml')
body('config.yaml 已加入 .gitignore，不會被 git 追蹤，需在各部署環境手動建立。')
doc.add_paragraph()
code_block(
    'mqtt:\n'
    '  broker: "109.123.238.225"    # MQTT Broker IP\n'
    '  port: 1883                   # MQTT port\n'
    '  username: ""                 # 帳號（無則留空）\n'
    '  password: ""                 # 密碼（無則留空）\n'
    '  client_id: "mqtt_sql_bridge" # MQTT client ID（需唯一）\n'
    '  keepalive: 60                # 心跳間隔（秒）\n'
    '  reconnect_delay: 10          # 斷線重連延遲（秒）\n'
    '  topics:\n'
    '    - topic: "zhongli/zone1/compressor/c2/telemetry"\n'
    '      table: "PROCESS_SGE125"  # 對應 DB table 名稱\n'
    '    - topic: "zhongli/zone2/compressor/c3/telemetry"\n'
    '      table: "PROCESS_BSAV55A"\n'
    '    - topic: "zhongli/zone2/compressor/c4/telemetry"\n'
    '      table: "PROCESS_SA37"\n'
    '\n'
    'db:\n'
    '  server: "YOUR_SQL_SERVER_IP"        # SQL Server IP\n'
    '  database: "YOUR_DATABASE_NAME"     # 資料庫名稱\n'
    '  username: "YOUR_DB_USERNAME"       # 帳號\n'
    '  password: "YOUR_DB_PASSWORD"       # 密碼\n'
    '  driver: "ODBC Driver 18 for SQL Server"\n'
    '  reconnect_delay: 5                # DB 斷線重連延遲（秒）\n'
    '\n'
    'log:\n'
    '  dir: "logs"\n'
    '  level: "INFO"   # DEBUG / INFO / WARNING / ERROR\n'
    '\n'
    'raw_json:\n'
    '  dir: "raw_json"\n'
    '\n'
    'queue_maxsize: 10000\n'
    'sampling_interval_seconds: 0  # 0=每筆寫入, N=每N秒平均一筆'
)

doc.add_paragraph()
h2('4.2 field_mapping.yaml')
body('定義 MQTT payload 的 JSON 欄位名稱對應到 DB 的欄位名稱（FIELD_1～FIELD_17）。')
doc.add_paragraph()
code_block(
    'FIELD_1:  compressor_drive_type\n'
    'FIELD_2:  area_entrance_instant_flow\n'
    'FIELD_3:  area_entrance_gas_pressure\n'
    '...（共 17 個欄位）\n'
    'FIELD_17: compressor_input_voltage'
)
body('新增/修改欄位對應後，需 rebuild container 才會生效。')

doc.add_paragraph()
h2('4.3 DB Table 欄位結構')
table2([
    ('CONTEXTID', 'NULL（保留）'),
    ('TIMETAG', 'MQTT payload 的 Timestamp（主鍵，ISO 8601 格式）'),
    ('TIME01', '寫入資料庫的當下時間（GETDATE()）'),
    ('TIME02', 'NULL（保留）'),
    ('TIME03', 'NULL（保留）'),
    ('FIELD_1 ～ FIELD_17', 'MQTT payload 感測器數值'),
], headers=['欄位', '說明'])

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 5. 日常操作
# ══════════════════════════════════════════════════════════
h1('5. 日常操作')

h2('5.1 確認服務是否正在運行')
body('在 WSL2 Ubuntu 終端機執行：')
code_block('docker ps')
body('正常輸出應可看到 mqtt_sql_bridge 的 container 狀態為 Up。')

doc.add_paragraph()
h2('5.2 查看即時 Log')
code_block(
    '# 即時追蹤（Ctrl+C 離開）\n'
    'docker logs mqtt_sql_bridge -f\n'
    '\n'
    '# 只看最新 50 行\n'
    'docker logs mqtt_sql_bridge --tail 50'
)
body('正常運行時 log 應顯示：')
bullet('[INFO] DB connected to ...')
bullet('[INFO] MQTT connected to ...')
bullet('[INFO] Subscribed: topic → table')

doc.add_paragraph()
h2('5.3 重啟 Container')
code_block('cd ~/mqtt-sql-bridge\ndocker compose restart')

doc.add_paragraph()
h2('5.4 修改設定後重啟')
body('修改 config.yaml 後（config 是 volume mount，不需 rebuild）：')
code_block('docker compose restart')
body('修改 field_mapping.yaml 或程式碼後（需 rebuild）：')
code_block('docker compose up --build -d')

doc.add_paragraph()
h2('5.5 停止 / 啟動')
code_block(
    '# 停止\n'
    'docker compose down\n'
    '\n'
    '# 啟動\n'
    'docker compose up -d'
)

doc.add_paragraph()
h2('5.6 開機自動啟動')
body('WSL2 已設定 /etc/wsl.conf 於啟動時自動執行 Docker。'
     'Windows 排程工作（WSL2-DockerAutoStart）會在登入時啟動 WSL2 並執行 docker compose up -d。')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 6. 常見問題排查
# ══════════════════════════════════════════════════════════
h1('6. 常見問題排查')

h2('6.1 DB 連線失敗')
code_block('[WARNING] DB reconnect failed: SSL Provider: certificate verify failed')
body('→ 確認 config.yaml 的 driver 為 "ODBC Driver 18 for SQL Server"，'
     'db_writer.py 連線字串包含 TrustServerCertificate=yes。')

doc.add_paragraph()
h2('6.2 DB 每隔 1~2 小時斷線')
code_block('[ERROR] DB insert error: Communication link failure')
body('→ 正常現象，系統會自動重連。防火牆切斷閒置 TCP 連線，'
     '已加入 5 分鐘心跳機制（SELECT 1）防止頻繁斷線。')

doc.add_paragraph()
h2('6.3 收到空 Payload')
code_block('[WARNING] Empty payload on topic: ..., skipped')
body('→ 正常，裝置偶爾會發空訊息，系統會自動跳過。')

doc.add_paragraph()
h2('6.4 PRIMARY KEY 重複')
code_block('[ERROR] DB insert error: 違反 PRIMARY KEY 條件約束')
body('→ 通常發生在 container 重啟後，MQTT broker 補發未確認的舊訊息。'
     'TIMETAG 相同的資料已存在 DB，不影響新資料寫入。')

doc.add_paragraph()
h2('6.5 WSL2 無法啟動')
code_block('錯誤碼: Wsl/Service/CreateInstance/E_FAIL')
body('→ 以系統管理員開啟 PowerShell 執行：')
code_block('wsl --shutdown\nwsl -d Ubuntu-22.04')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 7. 新增設備 Topic 流程
# ══════════════════════════════════════════════════════════
h1('7. 新增設備 / Topic 流程')

body('當需要新增一台設備的資料接收時，依序執行：')
doc.add_paragraph()

h3('步驟 1：在 config.yaml 新增 topic')
code_block(
    'topics:\n'
    '  - topic: "zhongli/zone1/compressor/c5/telemetry"\n'
    '    table: "PROCESS_NewDevice"   # DB table 名稱'
)

h3('步驟 2：確認 DB 有對應 table')
body('到 SQL Server 確認 PROCESS_NewDevice table 已建立，欄位包含：'
     'CONTEXTID, TIMETAG（PK）, TIME01, TIME02, TIME03, FIELD_1～FIELD_17。')

h3('步驟 3：重啟 container')
code_block('docker compose restart')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 8. 環境重建流程
# ══════════════════════════════════════════════════════════
h1('8. 環境重建流程（換機或重裝）')

body('若需要在新機器上重新部署，依序執行：')
doc.add_paragraph()

table2([
    ('1', '安裝 WSL2 + Ubuntu 22.04'),
    ('2', '在 Ubuntu 執行 curl -fsSL https://get.docker.com | sh 安裝 Docker Engine'),
    ('3', '執行 sudo usermod -aG docker $USER'),
    ('4', '設定開機自動啟動：echo -e "[boot]\\ncommand=\\"service docker start\\"" | sudo tee /etc/wsl.conf'),
    ('5', '複製專案至 WSL2：cp -r /mnt/c/Users/.../mqtt-sql-bridge ~/mqtt-sql-bridge'),
    ('6', '建立 config.yaml（參考第 4 章，此檔不在 git 內）'),
    ('7', '執行 cd ~/mqtt-sql-bridge && docker compose up --build -d'),
    ('8', '確認 docker logs mqtt_sql_bridge --tail 20 無錯誤'),
], headers=['步驟', '說明'])

# ══════════════════════════════════════════════════════════
# 儲存
# ══════════════════════════════════════════════════════════
out_path = r'C:\Users\autolab\Documents\code\mqtt-sql-bridge\MQTT_SQL_Bridge_交接文件.docx'
doc.save(out_path)
print(f'Done: {out_path}')
